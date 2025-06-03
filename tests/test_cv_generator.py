import os
import sys
import pytest
from docx.document import Document as DocumentClass

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
import cv_generator  # noqa: E402
from cv_generator import create_cv, main  # noqa: E402


def minimal_content():
    return {
        'personal_info': {'name': 'Name', 'contact': 'Contact'},
        'professional_summary': 'Summary.',
        'key_skills': ['Skill1', 'Skill2'],
        'accomplishments': [],
        'work_history': [],
        'education': [],
        'certifications': []
    }


def content_without_optional():
    return {
        'personal_info': {'name': 'Name', 'contact': 'Contact'},
        'professional_summary': 'Summary.',
        'accomplishments': [],
        'work_history': [],
        'education': [],
        'certifications': []
    }


def test_key_skills_section_present():
    doc = create_cv(minimal_content())
    texts = [p.text for p in doc.paragraphs]
    assert 'Key Skills' in texts
    index = texts.index('Key Skills')
    assert texts[index + 1] == 'Skill1'
    assert texts[index + 2] == 'Skill2'
    assert doc.paragraphs[index].style.name == 'CustomHeading'
    assert doc.paragraphs[index + 1].style.name == 'List Bullet'


def test_load_cv_content_returns_dict():
    content = cv_generator.load_cv_content('cv_content.json')
    assert isinstance(content, dict)


def test_load_cv_content_raises_for_missing_file(tmp_path):
    missing = tmp_path / "missing.json"
    with pytest.raises(FileNotFoundError):
        cv_generator.load_cv_content(missing)


def test_create_cv_returns_document_instance():
    content = cv_generator.load_cv_content('cv_content.json')
    doc = cv_generator.create_cv(content)
    assert isinstance(doc, DocumentClass)
    # Verify that at least the 'Professional Summary' heading is present
    texts = [p.text for p in doc.paragraphs]
    assert 'Professional Summary' in texts


def test_create_cv_handles_missing_optional_sections():
    doc = create_cv(content_without_optional())
    texts = [p.text for p in doc.paragraphs]
    assert 'Key Skills' not in texts


def test_cli_creates_output_file(tmp_path):
    output = tmp_path / "cv.docx"
    main(["-i", "cv_content.json", "-o", str(output)])
    assert output.exists()


def test_create_cv_with_template(tmp_path):
    template_path = tmp_path / "template.docx"
    # create a simple template with a placeholder paragraph
    from docx import Document
    template_doc = Document()
    template_doc.add_paragraph("Placeholder")
    template_doc.save(template_path)
    doc = create_cv(minimal_content(), template=str(template_path))
    texts = [p.text for p in doc.paragraphs]
    assert "Placeholder" in texts

