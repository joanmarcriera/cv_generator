"""Generate a CV in DOCX format from structured JSON content."""

import json
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def load_cv_content(filename):
    with open(filename, 'r') as file:
        return json.load(file)


def create_cv(content, template=None):
    """Create a CV Document.

    Parameters
    ----------
    content : dict
        Parsed JSON content describing the CV.
    template : str, optional
        Path to a .docx file to use as a starting template.
    """
    doc = Document(template) if template else Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Create a custom style for headings
    heading_style = doc.styles.add_style(
        'CustomHeading', WD_STYLE_TYPE.PARAGRAPH
    )
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(12)
    heading_style.font.bold = True
    paragraph_format = heading_style.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(4)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.add_run(content['personal_info']['name'] + '\n')
    run.font.size = Pt(14)
    run.bold = True
    run = header.add_run(content['personal_info']['contact'] + '\n')
    run.font.size = Pt(11)

    # Professional Summary
    doc.add_paragraph('Professional Summary', style='CustomHeading')
    doc.add_paragraph(content['professional_summary'])

    # Key Skills
    if content.get('key_skills'):
        doc.add_paragraph('Key Skills', style='CustomHeading')
        for skill in content['key_skills']:
            skill_para = doc.add_paragraph(skill, style='List Bullet')
            skill_para.paragraph_format.space_after = Pt(0)

    # Accomplishments
    doc.add_paragraph('Accomplishments', style='CustomHeading')
    for accomplishment in content['accomplishments']:
        doc.add_paragraph(accomplishment, style='List Bullet')

    # Work History
    doc.add_paragraph('Work History', style='CustomHeading')
    for position in content['work_history']:
        title_line = f"{position['title']} ({position['dates']})"
        doc.add_paragraph(title_line, style='CustomHeading')
        company = doc.add_paragraph(position['company'])
        company.style = 'Intense Quote'
        for responsibility in position['responsibilities']:
            para = doc.add_paragraph(responsibility, style='List Bullet')
            para.paragraph_format.space_after = Pt(0)

    # Education
    doc.add_paragraph('Education', style='CustomHeading')
    for edu in content['education']:
        para = doc.add_paragraph(edu, style='List Bullet')
        para.paragraph_format.space_after = Pt(0)

    # Certifications
    doc.add_paragraph('Certifications', style='CustomHeading')
    for cert in content['certifications']:
        para = doc.add_paragraph(cert, style='List Bullet')
        para.paragraph_format.space_after = Pt(0)

    return doc


def parse_args(args=None):
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description="Generate a CV from JSON data")
    parser.add_argument(
        "-i",
        "--input",
        default="cv_content.json",
        help="Path to JSON file with CV content",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="Joan_Marc_Riera_CV.docx",
        help="Output DOCX file name",
    )
    parser.add_argument(
        "-t",
        "--template",
        help="Optional DOCX template to use",
    )
    return parser.parse_args(args)


def main(args=None):
    options = parse_args(args)
    content = load_cv_content(options.input)
    doc = create_cv(content, template=options.template)
    doc.save(options.output)


if __name__ == "__main__":
    main()
